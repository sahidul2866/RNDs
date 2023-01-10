import os
import sys
from time import sleep

import pythoncom
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
from flask import send_file

import GraphMaker

# coverPageDict = {"EntityName": "Bangladesh", "BranchAuditMonthName": "Dhaka August", "Report#": "Report 123456",
#                  "AuditPerformer": "Md. Sahidul Islam", "AuditIssuer": "Pulok Bhai", "AuditDate": "21/03/2022"}
#
# sectionData = {"ContentOfEnvironment": "Environment Content", "ContentOfScope": "Scope Content",
#                "ContentOfOpinion": "Opinion Content"}
#
# description = r"Le Lorem Ipsum est simplement du faux texte employé dans la composition et la mise en page avant impression. Le Lorem Ipsum est le faux texte standard de l'imprimerie depuis les années 1500, quand un imprimeur anonyme assembla ensemble des morceaux de texte pour réaliser un livre spécimen de polices de texte. Il n'a pas fait que survivre cinq siècles, mais s'est aussi adapté à la bureautique informatique, sans que son contenu n'en soit modifié. Il a été popularisé dans les années 1960 grâce à la vente de feuilles Letraset contenant des passages du Lorem Ipsum, et, plus récemment, par son inclusion dans des applications de mise en page de texte, comme Aldus PageMaker."
#
# detailedIssueTableContent = {"IssueTitle": "Shezan Bhai", "IssueOwner": "Rahimin bhai", "IssueRating": "Five star",
#                              "IssueTargetDate": "23-03-2022", "IssueDescription": description}


def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application", pythoncom.CoInitialize())
    doc = word.Documents.Open(docx_file)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()


def update_texts(document, coverPageDict, sectionData, coverPage=True):
    # updating cover page #check for misuse
    for paragraph in document.paragraphs:
        if coverPage and paragraph.text.strip() in coverPageDict:
            if paragraph.text.strip() == "AuditDate":
                coverPage = False
            paragraph.text = coverPageDict[paragraph.text.strip()]
        # updating sections 1-3
        if paragraph.text.strip() in sectionData:
            paragraph.text = sectionData[paragraph.text.strip()]
            paragraph.style = document.styles['Normal']


def update_header(document, headerText, section=0):
    section = document.sections[section]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "\t\t" + headerText
    paragraph.style = document.styles["Header"]


def update_table(document, detailedIssueTableContent):
    table = document.tables[-1]
    table.cell(0, 1).text = detailedIssueTableContent["IssueTitle"]
    table.cell(1, 1).text = detailedIssueTableContent["IssueOwner"]
    table.cell(2, 1).text = detailedIssueTableContent["IssueRating"]
    table.cell(3, 1).text = detailedIssueTableContent["IssueTargetDate"]
    para = table.cell(4, 0).add_paragraph(
        "")  # = table.cell(4, 0).text.strip() + " " + detailedIssueTableContent["IssueDescription"]  #make issue description bold
    para.add_run("Description: ").bold = True
    para.add_run(detailedIssueTableContent["IssueDescription"])
    img = document.tables[-2].rows[0].cells[0].add_paragraph().add_run().add_picture('graph.png', width=Inches(5.5))


def page_break(paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_section(document, title, content):
    document.add_heading(title, 2)
    paragraph = document.add_paragraph(content)


def keep_table_on_one_page(doc):
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
        ppr = tag.get_or_add_pPr()
        ppr.keepNext_val = True


def callable_from_others(coverPageDict, sectionData, detailedIssueTableContent, graph):
    document = Document("INTERNAL AUDIT REPORT.docx")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    update_texts(document, coverPageDict, sectionData)
    update_header(document, coverPageDict["Report#"], 0)
    update_header(document, coverPageDict["Report#"], 1)
    GraphMaker.graphMaker(graph)
    update_table(document, detailedIssueTableContent)
    keep_table_on_one_page(document)


    document.save("template_output.docx")
    update_toc(r"C:\Users\Sahidul\Desktop\RNDs\WordGeneratorPy\template_output.docx")
    # openers = {'linux': 'libreoffice template_output.docx',
    #        'linux2': 'libreoffice template_output.docx',
    #        'darwin': 'open template_output.docx',
    #        'win32': 'start template_output.docx'}
    # os.system(openers[sys.platform])

    try:
        return send_file('template_output.docx', attachment_filename='output.docx')
    except Exception as e:
        return str(e)
