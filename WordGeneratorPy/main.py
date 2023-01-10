#First set directory where you want to save the file

import os
#os.chdir("D:/")

#Now import required packages

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

#Initialising document to make word file using python

document = Document()

#Code for making Table of Contents

paragraph = document.add_paragraph("My name is sahidul")
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p

#Giving headings that need to be included in Table of contents

document.add_heading("Network Connectivity")

paragraph = document.add_paragraph("My name is sahidul, hgfdhsuzxvchzx zcyc zsvuczxvcv csuydvyhc")

document.add_heading("Weather Stations")

#Saving the word file by giving name to the file

name = "mdh2"
document.save(name+".docx")

#Now check word file which got created

#Select "Right-click to update field text"
#Now right click and then select update field option
#and then click on update entire table

#Now,You will find Automatic Table of Contents