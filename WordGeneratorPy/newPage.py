import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding a paragraph
doc.add_heading('Page 1:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.')

# Adding a page break
doc.add_page_break()

# Adding a paragraph
doc.add_heading('Page 2:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.')

# Now save the document to a location
doc.save('gfg.docx')